import os
import platform
from transformers import AutoTokenizer, AutoModel
from docx import Document
import pandas as pd

tokenizer = AutoTokenizer.from_pretrained("/home/egnail/proj/ChatGLM2-6B-model", trust_remote_code=True)
model = AutoModel.from_pretrained("/home/egnail/proj/ChatGLM2-6B-model", trust_remote_code=True).cuda()
model = model.eval()
os_name = platform.system()
clear_command = 'cls' if os_name == 'Windows' else 'clear'
stop_stream = False

def main():
    dir = "/home/egnail/proj/NiCheng/泥城镇三重一大/泥委（2015）39号附件泥城镇关于“三重一大”事项集体决策制度的实施办法.docx"
    doc = Document(dir)
    key1 = ['一、基本要求', '（一）重大决策类：', '（二）重要干部任免、奖惩类：', '（三）重大项目类：', '（四）大额度资金使用类：']
    key2 = ['二、“三重一大”事项的具体内容及范围', '（二）重要干部任免、奖惩类：', '（三）重大项目类：', '（四）大额度资金使用类：', '三、主要程序']
    idx = []
    j = 0
    sub, str = [], []
    for i, page in enumerate(doc.paragraphs):
        text2 = page.text
        if key1[j] in text2:
            sub.append(i)
            j += 1
        if j == 5:
            break
    idx.append(sub)
    j = 0
    sub = []
    for i, page in enumerate(doc.paragraphs):
        text2 = page.text
        str.append(text2)
        if key2[j] in text2:
            sub.append(i)
            j += 1
        if j == 5:
            break
    idx.append(sub)
    first_ans = ('', '')
    q1 = ''.join(str[idx[0][0]: idx[1][0]])
    q2 = ''.join(str[idx[0][1]: idx[1][1]])
    q3 = ''.join(str[idx[0][2]: idx[1][2]])
    q4 = ''.join(str[idx[0][3]: idx[1][3]])
    q5 = ''.join(str[idx[0][4]: idx[1][4]])
    q6 = "根据前述的\"三重一大\"的具体内容及范围，你最后认为本次会议内容和议事结果属于哪一类？"
    q7 = "根据之前讲过的基本要求和\"三重一大\"的具体内容及范围，你简要总结本次会议内容和议事结果。"
    q8 = "根据之前讲过的基本要求和\"三重一大\"的具体内容及范围，你认为本次会议内容和议事结果合规吗？如不合规罗列出你的疑点"
    Ques = [q1, q2, q3, q4, q5, q6, q7, q8]
    query = "假设你是一位国家审计人员，现在需要你根据我给出的制度实施办法细则，对会议纪要的内容和结果进行回答。"
    Ques.insert(0, query)
    ############################################################################
    dir = "/home/egnail/proj/NiCheng/泥城镇三重一大/党委会镇长办公会会议纪要（泥城镇）.xlsx"
    memos = pd.read_excel(dir, sheet_name='镇长办公会', header=[1]).dropna()
    meet_cont = memos['会议内容']
    meet_result = memos['议事结果']
    past_key_values, history = None, []
    for i, _ in enumerate(Ques):
        current_length = 0
        if i < 6 and i > 0:
            query = "根据" + _ + "，你认为本次会议内容和议事结果属于哪一类？"
        else:
            query = _
        for response, history, past_key_values in model.stream_chat(tokenizer, query, history=history,
                                                                        past_key_values=past_key_values,
                                                                        return_past_key_values=True):
    
    
if __name__ == "__main__":
    main()