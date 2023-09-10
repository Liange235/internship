import pandas as pd
from pdf2image import convert_from_path, convert_from_bytes
import pytesseract
from pypdf import PdfReader
from docx import Document
from fnmatch import fnmatch
import difflib
import re
import cv2
import os
import numpy as np
from joblib import Parallel, delayed
from tqdm import tqdm
import datetime
starttime = datetime.datetime.now()

# paragraph = document.add_paragraph()
# run = paragraph.add_run(text2_filtered, style='Chinese')
# document.add_page_break()
# document.save("29、上会师报字（2021）第2856号-上海莱泽保险代理有限公司.docx")

def read_report(file, s1, s2):
    # flag = True
    key_str = 'Not Found'
    ins = 'Not Found'
    name = 'NotFound'
    sugg = 'Not Found'
    form = file.split('.')[-1]
    try:
        if form == 'pdf':
            reader = PdfReader(file)
            last_ind = 16 if len(reader.pages)>16 else len(reader.pages)
            images = convert_from_path(file, first_page=0, last_page=last_ind, dpi=600)
            for i, _ in enumerate(images):
                text2 = pytesseract.image_to_string(_, lang='chi_sim')  # ocr the img to text
                # text2_filtered = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text2)
                text2_filtered = text2.replace(' ', '')
                text2_filtered = re.sub(r'[\x00-\x09\x0b-\x1f\x7f-\x9f]', '', text2_filtered)
                text2_filtered = text2_filtered.replace(' ', '')
                text2_filtered = re.sub(r'\n{2,}', '\n', text2_filtered)
                if i == 0:
                    valid_s = []
                    for idx, str in enumerate(text2_filtered.split('\n')):
                        if str:
                            valid_s.append(str)
                    found_indices = [idx for idx, str in enumerate(valid_s) if '公司' in str]
                    found_indices = found_indices if found_indices else [1]
                    name = valid_s[found_indices[0]]
                    found_indices = [idx for idx, str in enumerate(valid_s) if '报告' in str]
                    found_indices = found_indices if found_indices else [3]
                    ins = valid_s[found_indices[0]+1]
                    if not re.match(r'\w+-\w+', ins):
                        text2 = pytesseract.image_to_string(images[1], lang='chi_sim')
                        text2_filtered = text2.replace(' ', '')
                        text2_filtered = re.sub(r'[\x00-\x09\x0b-\x1f\x7f-\x9f]', '', text2_filtered)
                        text2_filtered = text2_filtered.replace(' ', '')
                        text2_filtered = re.sub(r'\n{2,}', '\n', text2_filtered)
                        valid_s = []
                        for idx, str in enumerate(text2_filtered.split('\n')):
                            if str:
                                valid_s.append(str)
                        found_indices = [idx for idx, str in enumerate(valid_s) if re.match(r'.*以下简称.*', str)]
                        found_indices = found_indices if found_indices else [4]
                        ins = valid_s[found_indices[0]-2] + valid_s[found_indices[0]-1]
                elif fnmatch(text2_filtered, '*查结论*'):
                    text2_filtered = text2_filtered.replace('\n', '')
                    idx1 = text2_filtered.index(s1)
                    curr_id = text2_filtered[idx1-2]
                    key_str = text2_filtered[idx1:]
                    try:
                        id_id = s2.index(curr_id)
                        idx2 = text2_filtered.index(s2[id_id+1]+'、')
                        key_str = text2_filtered[idx1: idx2]
                        sugg = text2_filtered[idx2:]
                        if last_ind - i > 1:
                            next_image = images[i+1]
                            if not is_table(next_image) :
                                text2 = pytesseract.image_to_string(next_image, lang='chi_sim')
                                text2_filtered = text2.replace(' ', '')
                                text2_filtered = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text2_filtered)
                                text2_filtered = text2_filtered.replace(' ', '')
                                text2_filtered = re.sub(r'\n{2,}', '\n', text2_filtered)
                                sugg += text2_filtered
                    except ValueError:
                        if last_ind - i > 1:
                            next_image = images[i+1]
                            if not is_table(next_image) :
                                next_page = pytesseract.image_to_string(images[i+1], lang='chi_sim')
                                next_page = next_page.replace(' ', '')
                                NextPage_filtered = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', next_page)
                                NextPage_filtered = NextPage_filtered.replace(' ', '')
                                NextPage_filtered = re.sub(r'\n{2,}', '\n', NextPage_filtered)
                                if s2[id_id+1]+'、' in NextPage_filtered:
                                    idx2 = NextPage_filtered.index(s2[id_id+1]+'、')
                                    k2 = NextPage_filtered[:idx2]
                                    sugg = NextPage_filtered[idx2:]
                                    key_str += k2
                                else:
                                    key_str += NextPage_filtered
                    break
        elif form == 'docx':
            flag = True
            flag1 = False
            cont = []
            doc = Document(file)
            str = ''
            for i, page in enumerate(doc.paragraphs):
                text2 = page.text
                text2 = re.sub(r'\n{2,}', '\n', text2)
                if fnmatch(text2, '*被*查机构*') & flag:
                    name = text2.split('：')[-1].strip()[:-1]
                    idx = i
                    flag = False
                    continue
                elif fnmatch(text2, '*查结论*') | flag1:
                    if not flag1:
                        id_id = s2.index(text2[0])
                        str = s2[id_id+1]+'、'
                    cont.append(text2)
                    flag1 = True
                    if str in text2:
                        sugg = doc.paragraphs[i+1].text
                        flag1 = False
                        break
            text2 = doc.paragraphs[idx+1].text
            idx2 = text2.index('以下简称')
            ins = text2[:idx2-1]
            key_str = '\n'.join(cont[:-1])
    except Exception as e:
        print("Error: Exception occurred during processing.")
        print(e)
        print(file)
        return key_str, name, ins, sugg
    
    return key_str, name, ins, sugg

def is_table(image):
    image_array = np.array(image)
    ## 判断是否是表格
    gray = cv2.cvtColor(image_array, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=100, minLineLength=100, maxLineGap=10)
    if lines is not None and len(lines) > 4:
        return True
    else:
        return False

def get_files(dir, paths, form, fnames):
    for filepath, dirnames, filenames in os.walk(dir):
        for filename in filenames:
            if '核查报告' in filename:
                abs_p = os.path.join(filepath, filename)
                f = abs_p.split('.')
                paths.append(abs_p)
                form.append(f[1])
                fnames.append(filename)
    # 手动添加2022年的文件夹，由于从名称上无法有效判别是否是最终报告，因此先手动指定
    dir2 = dir + '/2022'
    for filepath, dirnames, filenames in os.walk(dir2):
        for filename in filenames:
            abs_p = os.path.join(filepath, filename)
            f = abs_p.split('.')
            paths.append(abs_p)
            form.append(f[1])
            fnames.append(filename)
    return paths, form, fnames
def my_fun(dir, s1, s2):
    con, nam, ins, sugg = read_report(dir, s1, s2)
    # print('Proc, '+ dir)
    # para.add_run(_)
    # para.add_run('\n')
    return con, nam, ins, sugg
def process_files(dir_sou, s1, s2):
    paths, form, fnames, dirs = [], [], [], []
    abs_path, form, file = get_files(dir_sou, paths, form, fnames)
    ##采用词向量匹配相似度，删除相似度>0.8的文件 鉴于可能是不同机构所出报告故先不处理
    # similarity = np.zeros((len(file), len(file)))
    # for i, _ in enumerate(file):
        # s1 = file[i].split('.')[0]
        # for j in range(i+1, len(file)):
            # s2 = file[j].split('.')[0]
            # similarity[i, j] = difflib.SequenceMatcher(autojunk=True, a=s1, b=s2).ratio()
    # idx = np.where(similarity>0.9)
    # del abs_path[idx[0]], form[idx[0]], file[idx[0]]
    for i, _ in enumerate(abs_path):
        str_block = _.split('/')
        dirs.append('\\'.join(_ for _ in str_block[4:]))
    # res = []
    # for _ in abs_path:
        # res.append(my_fun(_, s1, s2))
    # res.append(my_fun(dir, s1, s2))
    res = Parallel(n_jobs=16)(delayed(my_fun)(_, s1, s2) for _ in tqdm(abs_path))
    return res, form, file, dirs
######################################The main function######################################
def main():
    c_path = os.getcwd()
    dir = r"/mnt/nvme1n1/data/涉非核查（稳定处）/2020-2022律所核查报告(稳定处）"
    s1 = "检查结论"
    s2 = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    res, form, file, dirs = process_files(dir, s1, s2)
    KeyStr = [_[0] for _ in res]
    CompanyName = [_[1] for _ in res]
    Institution = [_[2] for _ in res]
    Suggestion = [_[3] for _ in res]
    df1 = pd.DataFrame({'机构': pd.Series(Institution), '企业': pd.Series(CompanyName), '源文件': pd.Series(file),
                        '格式': pd.Series(form), '相对路径': pd.Series(dirs), '核查结论': pd.Series(KeyStr), '核查建议': pd.Series(Suggestion)})
    # 将每个 DataFrame 导出到一个单独的工作表，并应用样式
    filepath = c_path+'/info_statistics2.xlsx'
    writer = pd.ExcelWriter(filepath, engine='openpyxl')
    df1.to_excel(writer, sheet_name='2020-2022律所核查报告(稳定处）')
    # df2.style.apply(highlight_row, idx=idx[1], axis=1).to_excel(writer, sheet_name='2020-2022会所核查报告（稳定处）')
    writer._save()

    endtime = datetime.datetime.now()
    print (f"Total training time: {(endtime - starttime).seconds:d}")
    
if __name__ == "__main__":
    main()